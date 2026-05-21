"""
utils/file_parser.py - File parsing utilities for PDF, DOCX, and TXT files

Uses PyMuPDF (fitz) for PDF and python-docx for Word documents.
"""

import fitz  # PyMuPDF
import docx
import logging

logger = logging.getLogger(__name__)


def parse_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using PyMuPDF.
    
    Args:
        file_path: Path to the PDF file
    
    Returns:
        Extracted text as a string
    """
    try:
        text_parts = []
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc):
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"Parsed PDF: {len(full_text)} characters from {len(text_parts)} pages")
        return full_text

    except Exception as e:
        logger.error(f"PDF parsing error: {e}")
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def parse_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.
    
    Args:
        file_path: Path to the DOCX file
    
    Returns:
        Extracted text as a string
    """
    try:
        doc = docx.Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        full_text = "\n\n".join(paragraphs)
        logger.info(f"Parsed DOCX: {len(full_text)} characters")
        return full_text

    except Exception as e:
        logger.error(f"DOCX parsing error: {e}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def parse_text_file(file_path: str) -> str:
    """
    Read a plain text file.
    
    Args:
        file_path: Path to the text file
    
    Returns:
        File content as a string
    """
    try:
        # Try UTF-8 first, fall back to latin-1
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()

    except Exception as e:
        logger.error(f"Text file parsing error: {e}")
        raise ValueError(f"Failed to read text file: {str(e)}")


def get_file_type(filename: str) -> str:
    """
    Determine file type from extension.
    
    Returns:
        'pdf', 'docx', 'txt', or raises ValueError
    """
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    
    type_map = {
        "pdf": "pdf",
        "docx": "docx",
        "doc": "docx",
        "txt": "txt",
        "text": "txt",
        "md": "txt",
    }
    
    if ext not in type_map:
        raise ValueError(f"Unsupported file type: .{ext}. Supported: PDF, DOCX, TXT")
    
    return type_map[ext]
